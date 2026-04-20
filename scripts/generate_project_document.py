"""
Generates Greenwood_Zoo_Project_Document.docx in the project root.
Run: python scripts/generate_project_document.py
Requires: pip install python-docx
"""
from pathlib import Path

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_LINE_SPACING

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "Greenwood_Zoo_Project_Document.docx"


def add_body(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    for run in p.runs:
        run.font.size = Pt(11)


def main() -> None:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    doc.add_heading("Greenwood Wildlife Zoo — Project Document", 0)
    add_body(
        doc,
        "Team 9, COSC 3380. This document describes the main functions of the zoo management website: "
        "who uses each part, what each page or script does, and how the pieces fit together. "
        "File names refer to the webapp folder unless noted otherwise.",
    )

    sections = [
        (
            "1. Purpose of the system",
            [
                "The site supports two audiences: customers (tickets, gift shop, restaurant, cart, checkout) "
                "and staff (animals, employees, health, sales, alerts, reports). Staff sign in with a username; "
                "customers can use the same login form with their email address (see login_handler.php). "
                "PHP sessions store user_id and role for staff, or customer_id for customers.",
            ],
        ),
        (
            "2. Public marketing site",
            [
                "index.php — Public home page with zoo information, hours, and links to login, signup, and animals.php.",
                "index.css — Styling used mainly by the customer-facing layout and customer dashboard.",
                "animals.php — Public-facing animals listing for visitors who are not logged in.",
                "animals/elephants.php — Example or themed animal subpage under the public animals area.",
                "login.html — Staff login form (posts to login_handler.php).",
                "sign-in.html and unified_login.php — Alternate or combined entry points if your deployment links to them.",
                "signup.html and signup.php — New customer registration.",
                "customer-login.html and customer_login.php — Customer-only sign-in entry points.",
            ],
        ),
        (
            "3. Authentication and session",
            [
                "login_handler.php — Accepts POST login and password. If the value matches a customer email, "
                "it starts a customer session and sends the user to customer-dashboard.php. If it matches a "
                "systemuser username, it starts a staff session and redirects by role: admin to dashboard.php, "
                "caretaker to caretaker_dashboard.php, vet to vet_dashboard.php, and all other staff roles to dashboard.php.",
                "session_bootstrap.php — Included at the top of most PHP pages to start the session safely.",
                "logout.php — Clears the session and logs the user out.",
                "change-password.php — Lets a logged-in staff user change their password.",
            ],
        ),
        (
            "4. Customer area",
            [
                "customer-dashboard.php — Home after customer login. Shows welcome content, upcoming visit count, "
                "shortcuts to buy tickets, ticket history, restaurant, gift shop, and featured animals.",
                "customer_nav.php — Shared navigation for customers: home, dashboard, buy tickets, gift shop, "
                "restaurant, animals report, profile, cart, logout.",
                "buy_tickets.php — Lets customers pick ticket types and visit dates and add tickets to the cart.",
                "customer_tickets_report.php — Lists the signed-in customer’s ticket-related orders or history.",
                "giftshop.php — Browsable gift shop catalog; customers add items to cart. Staff may use a preview query string.",
                "restaurant.php — Restaurant menu by stall; customers add food items to the cart.",
                "cart.php — Shows cart contents (tickets, shop, food) and links toward checkout.",
                "cart_action.php — Backend handler for add, update, or remove cart lines via POST or GET; redirects back to cart or a given page.",
                "checkout.php — Completes purchase from the cart and creates orders in the database.",
                "customer_profile.php — View or edit customer profile information.",
                "customer_animals_report.php — Customer-facing list or search of animals in the collection.",
            ],
        ),
        (
            "5. Main staff dashboard (dashboard.php)",
            [
                "Most roles except caretaker and vet land here after login. The page title and tiles change by role.",
                "Admin sees summary counts: total animals, employees, health alerts, gift shop alerts, restaurant alerts, "
                "and today’s revenue from daily_revenue. Banners link to health-reports.php, shop_alerts.php, and restaurant_alerts.php.",
                "Animals and enclosures: add-animal.php (admin or caretaker), animals_report.php, health-reports.php (admin or vet).",
                "Staff management (admin): add-employee.php, employees_report.php.",
                "Revenue and tickets (admin or cashier): add-ticket.php. Admin also gets revenue_report.php (full revenue breakdown).",
                "Admin gift shop and restaurant shortcuts: add-gift-shop-item.php, shop_alerts.php, add-restaurant-item.php, restaurant_alerts.php.",
                "Gift Shop Employee section: restock banner, monthly snapshot, tiles for add-gift-shop-item.php, giftshop preview, "
                "add-order.php (record a sale), shop_alerts.php.",
                "Restaurant Employee section: restock banner, monthly snapshot, tiles for add-restaurant-item.php, add-restaurant-order.php, "
                "restaurant.php menu view, restaurant_alerts.php.",
                "Header actions: change-password.php, logout.php. Admin also includes admin_header_cart_profile.inc.php for cart and profile links.",
            ],
        ),
        (
            "6. Caretaker dashboard (caretaker_dashboard.php)",
            [
                "Focused workspace for caretakers: links back to main staff dashboard or vet dashboard, change password, logout.",
                "Tiles: add-animal.php, animals_report.php, health-reports.php, and an on-page care table anchor for daily tasks.",
            ],
        ),
        (
            "7. Veterinarian dashboard (vet_dashboard.php)",
            [
                "Vet landing page with staff dashboard link, change password, logout.",
                "Tiles: animals_report.php, health-reports.php, caretaker care table (caretaker_dashboard.php#care-table), "
                "and a default vetanimalupdate.php link for the next animal needing updates.",
                "Lists animals scheduled for vet attention with buttons to vetanimalupdate.php?id=… for each animal.",
            ],
        ),
        (
            "8. Animal and health record pages",
            [
                "add-animal.php — Form to register a new animal in the database.",
                "edit_animal.php — Edit an existing animal record.",
                "delete_animal.php — Remove an animal record (subject to app rules).",
                "animals_report.php — Searchable or filterable list of animals for staff.",
                "health-reports.php — Health status and medical reporting for staff and vets.",
                "vetanimalupdate.php — Detailed vet visit or health update form for one animal (by id).",
                "vetnotification.php — Vet-related notifications or messaging as implemented.",
                "vet_alerts_real_time.php — Live or polled vet alerts view if used in your deployment.",
            ],
        ),
        (
            "9. Employee management",
            [
                "add-employee.php — Register a new employee and system user.",
                "employees_report.php — List or filter employees.",
                "edit_employee.php — Edit employee details.",
                "delete_employee.php — Remove or deactivate an employee per business rules.",
            ],
        ),
        (
            "10. Tickets (staff)",
            [
                "add-ticket.php — Create or configure ticket products for sale.",
                "edit_tickets.php — Edit ticket definitions.",
                "delete_tickets.php — Remove ticket types.",
                "tickets_report.php — Staff report of ticket-related data.",
            ],
        ),
        (
            "11. Revenue and sales reports (mostly admin)",
            [
                "revenue_report.php — Consolidated revenue report (tickets, food, shop) with filters and tabs. Admin only.",
                "sales_report.php — Detailed gift shop sales with filters and charts. Admin only.",
                "restaurant_sales_report.php — Detailed restaurant sales. Admin only.",
            ],
        ),
        (
            "12. Gift shop operations",
            [
                "add-gift-shop-item.php — Add new shop products and stock.",
                "add-order.php — Staff records an in-person gift shop sale against inventory.",
                "shop_alerts.php — View and resolve low-stock or restock alerts for the gift shop.",
                "apply_shop_stock_trigger.php — Helper to install or verify database trigger behavior for shop stock.",
                "api_gift_shop_monthly_top.php — JSON API for top-selling shop lines in the current month (used for dashboard charts; "
                "allowed for admin and Gift Shop Employee).",
                "create_gift_shop_user.php — Utility script to create a gift shop role user for testing or setup.",
            ],
        ),
        (
            "13. Restaurant operations",
            [
                "add-restaurant-item.php — Add menu items and stock for a stall.",
                "add-restaurant-order.php — Staff records a restaurant sale.",
                "restaurant_alerts.php — View and resolve restaurant restock alerts.",
            ],
        ),
        (
            "14. Staff account and misc pages",
            [
                "staff_home.php — PHP helpers only: detects vet-like roles and returns the correct home URL "
                "(vet_dashboard.php, caretaker_dashboard.php, or dashboard.php) for links and redirects.",
                "staff_account.php — Admin-only profile page showing username, role, and employee details from the database.",
                "admin_header_cart_profile.inc.php — Shared header fragment: cart, profile, logout for admin on many PHP pages.",
                "style.css — Main stylesheet for staff and many internal pages.",
                "validation.js — Client-side validation used by selected forms.",
            ],
        ),
        (
            "15. Database and configuration",
            [
                "db.php — Creates the PDO connection to MySQL (host, database name, credentials, SSL options for Azure).",
                "DigiCertGlobalRootCA.crt.pem — Certificate bundle used when connecting to Azure Database for MySQL with SSL.",
            ],
        ),
        (
            "16. Setup and maintenance scripts (not for public production servers)",
            [
                "create-admin.php — Inserts a default admin employee and system user if missing (for first-time local setup).",
                "create_test_users.php — Inserts sample caretaker or vet users for demos.",
                "test.php — Simple PHP environment check.",
            ],
        ),
        (
            "17. Repository folders outside webapp",
            [
                "sql/trigger_shop_stock_alert.sql — Optional MySQL trigger for gift shop stock alerts.",
                ".github/workflows/main_team9zooproject.yml — GitHub Action to build and deploy the PHP app to Azure on push to main.",
            ],
        ),
        (
            "18. How roles reach different screens",
            [
                "Caretaker and vet are redirected to their own dashboards from login_handler.php. "
                "Admin, cashier, gift shop employee, and restaurant employee use dashboard.php, which uses exact role strings "
                "such as admin, Gift Shop Employee, and Restaurant Employee for some checks. "
                "Customers never use dashboard.php; they use customer-dashboard.php and related pages.",
            ],
        ),
    ]

    for title, paragraphs in sections:
        doc.add_heading(title, level=1)
        for para in paragraphs:
            add_body(doc, para)

    doc.add_page_break()
    doc.add_heading("End of document", level=1)
    add_body(
        doc,
        "Generated from the project codebase. If the team adds new PHP pages, re-run scripts/generate_project_document.py "
        "after updating this script, or edit the Word file by hand.",
    )

    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    main()
